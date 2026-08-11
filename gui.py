"""
ProtocolBot - Графический интерфейс для создания протоколов об обучении
Запуск: python gui.py
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.simpledialog import askstring

import os
import re
from datetime import datetime
from pathlib import Path
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from permits import PERMITS, parse_permits, SHEET_GROUPS, normalize_permit_codes
from create_protocol import extract_from_excel, extract_from_docx, extract_from_doc, extract_excel_tables
from config_loader import get_config, save_config

# Логи: все ошибки пишутся в protocolbot.log рядом с программой.
import app_log
app_log.install()


def split_permit_groups(codes_str):
    """
    Разбирает строку допусков с привязкой к группам вида:
        "2: 8,10,11; 3: 5,6,7"
    Возвращает словарь {группа: строка_кодов}.
    Старый формат без групп ("8,10,11") вернётся как {"": "8,10,11"}.
    """
    result = {}
    s = str(codes_str or "").strip()
    if not s:
        return result

    if ":" not in s:
        return {"": s}

    for chunk in s.split(";"):
        chunk = chunk.strip()
        if not chunk:
            continue
        if ":" in chunk:
            g, _, codes = chunk.partition(":")
            g = g.strip()
            codes = codes.strip()
            if g and codes:
                result[g] = codes
        else:
            result[""] = chunk
    return result


# =====================================================================
# Окно настроек распознавания (OCR).
# =====================================================================
class SettingsDialog(tk.Toplevel):
    def __init__(self, parent, config, on_saved=None):
        super().__init__(parent)
        self.title("Настройки распознавания (OCR)")
        self.geometry("760x360")
        self.transient(parent)
        self.grab_set()

        self.on_saved = on_saved
        self.config = dict(config)

        self.tess_var = tk.StringVar(value=self.config.get("tesseract_path", ""))
        self.pop_var = tk.StringVar(value=self.config.get("poppler_path", ""))
        self.lang_var = tk.StringVar(value=self.config.get("ocr_lang", "rus"))
        self.dpi_var = tk.StringVar(value=str(self.config.get("ocr_dpi", 300)))

        pad = {"padx": 10, "pady": 6}

        tk.Label(
            self,
            text="Эти настройки сохраняются в config.json и нужны, чтобы программа\n"
                 "работала на любом компьютере без правки кода.",
            justify=tk.LEFT,
            fg="#555555"
        ).pack(anchor=tk.W, **pad)

        body = ttk.Frame(self)
        body.pack(fill=tk.X, **pad)

        ttk.Label(body, text="Tesseract (tesseract.exe):").grid(row=0, column=0, sticky=tk.W, pady=4)
        ttk.Entry(body, textvariable=self.tess_var, width=58).grid(row=0, column=1, padx=5, pady=4)
        ttk.Button(body, text="Обзор…", command=self._browse_tess).grid(row=0, column=2, pady=4)
        self.tess_status = tk.Label(body, text="", width=14, anchor=tk.W)
        self.tess_status.grid(row=0, column=3, padx=5, pady=4)

        ttk.Label(body, text="Poppler (папка bin):").grid(row=1, column=0, sticky=tk.W, pady=4)
        ttk.Entry(body, textvariable=self.pop_var, width=58).grid(row=1, column=1, padx=5, pady=4)
        ttk.Button(body, text="Обзор…", command=self._browse_pop).grid(row=1, column=2, pady=4)
        self.pop_status = tk.Label(body, text="", width=14, anchor=tk.W)
        self.pop_status.grid(row=1, column=3, padx=5, pady=4)

        ttk.Label(body, text="Язык распознавания:").grid(row=2, column=0, sticky=tk.W, pady=4)
        ttk.Combobox(
            body, textvariable=self.lang_var, values=["rus", "rus+eng"],
            width=12, state="readonly"
        ).grid(row=2, column=1, sticky=tk.W, padx=5, pady=4)
        ttk.Label(
            body,
            text="«rus» — чисто по-русски; «rus+eng» — если много латиницы/марок",
            foreground="#888888"
        ).grid(row=2, column=2, columnspan=2, sticky=tk.W, padx=5)

        ttk.Label(body, text="Качество скана (dpi):").grid(row=3, column=0, sticky=tk.W, pady=4)
        ttk.Combobox(
            body, textvariable=self.dpi_var, values=["200", "300", "400", "600"],
            width=12
        ).grid(row=3, column=1, sticky=tk.W, padx=5, pady=4)
        ttk.Label(
            body,
            text="больше = чётче, но медленнее (обычно хватает 300)",
            foreground="#888888"
        ).grid(row=3, column=2, columnspan=2, sticky=tk.W, padx=5)

        self.tess_var.trace_add("write", lambda *a: self._refresh_status())
        self.pop_var.trace_add("write", lambda *a: self._refresh_status())
        self._refresh_status()

        btn = ttk.Frame(self)
        btn.pack(fill=tk.X, padx=10, pady=12)
        ttk.Button(btn, text="Сохранить", command=self._save).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn, text="Отмена", command=self.destroy).pack(side=tk.RIGHT)

        self.protocol("WM_DELETE_WINDOW", self.destroy)
        parent.wait_window(self)

    def _browse_tess(self):
        path = filedialog.askopenfilename(
            title="Выберите tesseract.exe",
            filetypes=[("Tesseract", "tesseract.exe"), ("Все файлы", "*.*")],
            initialvalue=self.tess_var.get()
        )
        if path:
            self.tess_var.set(path)

    def _browse_pop(self):
        path = filedialog.askdirectory(
            title="Выберите папку bin утилит Poppler",
            initialdir=self.pop_var.get() or None
        )
        if path:
            self.pop_var.set(path)

    def _refresh_status(self):
        tess = self.tess_var.get().strip()
        if tess and os.path.isfile(tess):
            self.tess_status.config(text="✓ найден", fg="#2e7d32")
        elif tess:
            self.tess_status.config(text="✗ не найден", fg="#c62828")
        else:
            self.tess_status.config(text="— не задан", fg="#c62828")

        pop = self.pop_var.get().strip()
        if pop and os.path.isdir(pop):
            self.pop_status.config(text="✓ найден", fg="#2e7d32")
        elif pop:
            self.pop_status.config(text="✗ не найден", fg="#c62828")
        else:
            self.pop_status.config(text="— не задан", fg="#c62828")

    def _save(self):
        try:
            dpi = int(self.dpi_var.get().strip() or 300)
        except ValueError:
            messagebox.showerror("Ошибка", "Качество скана (dpi) должно быть числом")
            return

        new_config = {
            "tesseract_path": self.tess_var.get().strip(),
            "poppler_path": self.pop_var.get().strip(),
            "ocr_lang": self.lang_var.get().strip() or "rus",
            "ocr_dpi": dpi,
        }

        save_config(new_config)
        self.config = new_config

        if self.on_saved:
            self.on_saved()

        self._refresh_status()
        messagebox.showinfo(
            "Сохранено",
            "Настройки сохранены в config.json и применены.\n"
            "Перезапуск программы не требуется."
        )
        self.destroy()


# =====================================================================
# Окно допусков: три столбца как в листе регистрации,
# галочки хранятся РАЗДЕЛЬНО по группам.
# Результат - строка вида "2: 8,10,11; 3: 5,6,7".
# =====================================================================
class PermitEditorDialog(tk.Toplevel):
    def __init__(self, parent, current_codes="", title="Допуски", default_groups=None):
        super().__init__(parent)
        self.title(title)
        self.geometry("1150x760")
        self.transient(parent)
        self.grab_set()

        self.result = None
        self.default_groups = default_groups or ["1", "2", "3"]
        self.code_vars = {}  # ключ: (группа, код)

        groups_map = split_permit_groups(current_codes)
        # Старый формат без групп раскидываем по группам человека.
        if "" in groups_map:
            flat = groups_map.pop("")
            for g in self.default_groups:
                groups_map.setdefault(g, flat)

        top = ttk.Frame(self, padding=10)
        top.pack(fill=tk.X)

        ttk.Label(top, text="Коды:").pack(side=tk.LEFT)

        self.entry = ttk.Entry(top)
        self.entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        if current_codes:
            self.entry.insert(0, str(current_codes))

        self.entry.bind("<KeyRelease>", lambda e: self._sync_checks_from_entry())

        self.warn = ttk.Label(self, text="", foreground="red")
        self.warn.pack(fill=tk.X, padx=10)

        mid = ttk.Frame(self)
        mid.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        for col_idx, (group_key, col_title) in enumerate((
            ("1", "Виды работ (для 1 группы)"),
            ("2", "Виды работ (для 2 группы)"),
            ("3", "Виды работ (для 3 группы)"),
        )):
            col_frame = ttk.Frame(mid, padding=6)
            col_frame.grid(row=0, column=col_idx, sticky="nsew", padx=4)
            mid.columnconfigure(col_idx, weight=1, uniform="cols")
            mid.rowconfigure(0, weight=1)

            ttk.Label(
                col_frame,
                text=col_title,
                font=("Arial", 10, "bold")
            ).pack(anchor=tk.W, pady=(0, 4))

            for code in SHEET_GROUPS[group_key]:
                self._add_checkbox(col_frame, group_key, code, groups_map)

        buttons = ttk.Frame(self, padding=10)
        buttons.pack(fill=tk.X)

        ttk.Button(buttons, text="OK", command=self._ok).pack(side=tk.RIGHT, padx=5)
        ttk.Button(buttons, text="Отмена", command=self._cancel).pack(side=tk.RIGHT)

        self.protocol("WM_DELETE_WINDOW", self._cancel)

        parent.wait_window(self)

    def _add_checkbox(self, parent, group, code, groups_map):
        selected, _ = normalize_permit_codes(groups_map.get(group, ""))
        var = tk.BooleanVar(value=code in selected)
        self.code_vars[(group, code)] = var

        text = str(PERMITS[code]).strip()

        tk.Checkbutton(
            parent,
            text=f"{code} — {text}",
            variable=var,
            command=self._sync_entry_from_checks,
            wraplength=320,
            anchor=tk.W,
            justify=tk.LEFT,
            bd=0,
            highlightthickness=0
        ).pack(anchor=tk.W, pady=1)

    def _sync_entry_from_checks(self):
        by_group = {}
        for (group, code), var in self.code_vars.items():
            if var.get():
                by_group.setdefault(group, []).append(code)

        parts = []
        for group in ("1", "2", "3"):
            if by_group.get(group):
                parts.append(f"{group}: " + ",".join(map(str, sorted(by_group[group]))))

        self.entry.delete(0, tk.END)
        self.entry.insert(0, "; ".join(parts))
        self._update_warning()

    def _sync_checks_from_entry(self):
        groups_map = split_permit_groups(self.entry.get())

        checked = set()
        invalid = []
        for group, codes in groups_map.items():
            valid, inv = normalize_permit_codes(codes)
            invalid += inv
            for code in valid:
                checked.add((group, code))

        for key, var in self.code_vars.items():
            var.set(key in checked)

        self._update_warning(invalid)

    def _update_warning(self, invalid=None):
        if invalid is None:
            invalid = []
            for codes in split_permit_groups(self.entry.get()).values():
                _, inv = normalize_permit_codes(codes)
                invalid += inv
        if invalid:
            self.warn.config(text="Неизвестные коды: " + ", ".join(invalid))
        else:
            self.warn.config(text="")

    def _ok(self):
        self._sync_entry_from_checks()
        self.result = self.entry.get().strip()
        self.destroy()

    def _cancel(self):
        self.result = None
        self.destroy()


# =====================================================================
# Главное окно программы.
# =====================================================================
class ProtocolBotApp:
    def __init__(self, root):
        self.root = root
        self.root.title("ProtocolBot - Создание протоколов")
        self.root.geometry("1000x800")
        self.root.minsize(900, 700)

        self.bg_color = "#f5f5f5"
        self.accent_color = "#2196F3"
        self.success_color = "#4CAF50"
        self.text_color = "#333333"

        self.root.configure(bg=self.bg_color)

        self.file_path = None
        self.participants = []
        self.permits_data = {}
        self.protocol_info = {}

        self.config = get_config()

        self.create_widgets()

    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # === ЗАГОЛОВОК ===
        header_frame = tk.Frame(main_frame, bg=self.accent_color, height=60)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        header_frame.pack_propagate(False)

        tk.Label(
            header_frame,
            text="ProtocolBot",
            font=("Arial", 20, "bold"),
            fg="white",
            bg=self.accent_color
        ).pack(side=tk.LEFT, padx=20)

        tk.Label(
            header_frame,
            text="Создание протоколов об обучении",
            font=("Arial", 12),
            fg="white",
            bg=self.accent_color
        ).pack(side=tk.LEFT, padx=10)

        # === ШАГ 1: ВЫБОР ФАЙЛА ===
        step1_frame = ttk.LabelFrame(main_frame, text="Шаг 1: Выберите файл со списком", padding=15)
        step1_frame.pack(fill=tk.X, pady=(0, 15))

        file_frame = ttk.Frame(step1_frame)
        file_frame.pack(fill=tk.X)

        self.file_label = ttk.Label(file_frame, text="Файл не выбран", foreground="gray")
        self.file_label.pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(
            file_frame,
            text="Выбрать файл",
            command=self.select_file
        ).pack(side=tk.RIGHT)

        pages_frame = ttk.Frame(step1_frame)
        pages_frame.pack(fill=tk.X, pady=(8, 0))

        ttk.Label(pages_frame, text="Страницы PDF:", width=20).pack(side=tk.LEFT)
        self.pages_var = tk.StringVar(value=str(self.config.get("pdf_pages", "") or ""))
        self.pages_entry = ttk.Entry(pages_frame, textvariable=self.pages_var, width=12)
        self.pages_entry.pack(side=tk.LEFT, padx=5)
        ttk.Label(
            pages_frame,
            text="например: 1 или 1-3; пусто = все страницы (значение запоминается)",
            foreground="gray"
        ).pack(side=tk.LEFT, padx=5)

        # === ШАГ 2: ДАННЫЕ ПРОТОКОЛА ===
        step2_frame = ttk.LabelFrame(main_frame, text="Шаг 2: Данные протокола", padding=15)
        step2_frame.pack(fill=tk.X, pady=(0, 15))

        row1 = ttk.Frame(step2_frame)
        row1.pack(fill=tk.X, pady=2)

        ttk.Label(row1, text="Номер протокола:", width=20).pack(side=tk.LEFT)
        self.protocol_num = ttk.Entry(row1, width=30)
        self.protocol_num.pack(side=tk.LEFT, padx=5)

        ttk.Label(row1, text="Дата:", width=10).pack(side=tk.LEFT, padx=(20, 0))
        self.date_entry = ttk.Entry(row1, width=15)
        self.date_entry.insert(0, datetime.now().strftime("%d.%m.%Y"))
        self.date_entry.pack(side=tk.LEFT, padx=5)

        row2 = ttk.Frame(step2_frame)
        row2.pack(fill=tk.X, pady=2)

        ttk.Label(row2, text="Группа по умолчанию:", width=20).pack(side=tk.LEFT)

        self.group_var = tk.StringVar(value="3")

        ttk.Radiobutton(row2, text="1 группа", variable=self.group_var, value="1").pack(side=tk.LEFT)
        ttk.Radiobutton(row2, text="2 группа", variable=self.group_var, value="2").pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(row2, text="3 группа", variable=self.group_var, value="3").pack(side=tk.LEFT)

        ttk.Label(row2, text="Часы:", width=10).pack(side=tk.LEFT, padx=(20, 0))
        self.hours_var = tk.StringVar(value="32")
        ttk.Combobox(
            row2,
            textvariable=self.hours_var,
            values=["16", "24", "32", "40", "72"],
            width=10
        ).pack(side=tk.LEFT, padx=5)

        row3 = ttk.Frame(step2_frame)
        row3.pack(fill=tk.X, pady=2)

        ttk.Label(row3, text="Организация:", width=20).pack(side=tk.LEFT)
        self.org_var = tk.StringVar(value="")
        self.org_entry = ttk.Entry(row3, textvariable=self.org_var)
        self.org_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        ttk.Label(
            step2_frame,
            text="Организация-заказчик (одна на всю заявку) попадёт во 2-й столбец под ФИО, "
                 "если у человека нет подразделения.",
            foreground="gray",
            wraplength=900,
            justify=tk.LEFT
        ).pack(anchor=tk.W, pady=(5, 0))

        ttk.Label(
            step2_frame,
            text="Если групп несколько, программа спросит: отдельные протоколы (Транснефть) "
                 "или один общий (человек с несколькими группами получит по строке на группу).",
            foreground="gray"
        ).pack(anchor=tk.W, pady=(2, 0))

        # === ШАГ 3: УЧАСТНИКИ ===
        step3_frame = ttk.LabelFrame(main_frame, text="Шаг 3: Участники, группы и допуски", padding=15)
        step3_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))

        toolbar = ttk.Frame(step3_frame)
        toolbar.pack(fill=tk.X, pady=(0, 8))

        ttk.Button(
            toolbar,
            text="＋ Добавить участника",
            command=self.add_participant
        ).pack(side=tk.LEFT, padx=(0, 8))

        ttk.Button(
            toolbar,
            text="− Удалить выбранных",
            command=self.remove_participants
        ).pack(side=tk.LEFT)

        columns = ("num", "name", "position", "subdivision", "group", "permits")

        self.tree = ttk.Treeview(
            step3_frame,
            columns=columns,
            show="headings",
            height=10,
            selectmode="extended"
        )

        self.tree.heading("num", text="№")
        self.tree.heading("name", text="ФИО")
        self.tree.heading("position", text="Должность")
        self.tree.heading("subdivision", text="Подразделение")
        self.tree.heading("group", text="Группа")
        self.tree.heading("permits", text="Коды допусков")

        self.tree.column("num", width=40, anchor=tk.CENTER)
        self.tree.column("name", width=210)
        self.tree.column("position", width=160)
        self.tree.column("subdivision", width=170)
        self.tree.column("group", width=70, anchor=tk.CENTER)
        self.tree.column("permits", width=140, anchor=tk.CENTER)

        scrollbar = ttk.Scrollbar(step3_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)

        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.tree.bind("<Double-1>", self.on_tree_double_click)
        self.tree.bind("<Button-3>", self.on_tree_right_click)

        # === ПОДСКАЗКА + НАСТРОЙКИ ===
        hint_frame = ttk.Frame(main_frame)
        hint_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(
            hint_frame,
            text="Двойной клик — редактировать ячейку. Правый клик — массовые действия. Коды: 2: 8,10; 3: 5,6",
            foreground="gray"
        ).pack(side=tk.LEFT)

        ttk.Button(
            hint_frame,
            text="Справка по допускам",
            command=self.show_permits_help
        ).pack(side=tk.RIGHT)

        ttk.Button(
            hint_frame,
            text="⚙ Настройки OCR",
            command=self.open_settings
        ).pack(side=tk.RIGHT, padx=(0, 8))

        # === КНОПКИ ===
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill=tk.X)

        ttk.Button(
            buttons_frame,
            text="Сгенерировать протокол",
            command=self.generate_protocol
        ).pack(side=tk.RIGHT, padx=5)

        ttk.Button(
            buttons_frame,
            text="Очистить",
            command=self.clear_all
        ).pack(side=tk.RIGHT, padx=5)

        self.status_var = tk.StringVar(value="Готово к работе")
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(fill=tk.X, pady=(10, 0))

    # -----------------------------------------------------------------
    # Настройки OCR
    # -----------------------------------------------------------------
    def open_settings(self):
        SettingsDialog(self.root, self.config, on_saved=self._on_config_saved)

    def _on_config_saved(self):
        self.config = get_config()

    def _check_ocr_paths(self):
        tess = (self.config.get("tesseract_path") or "").strip()
        pop = (self.config.get("poppler_path") or "").strip()

        tess_ok = bool(tess) and os.path.isfile(tess)
        pop_ok = bool(pop) and os.path.isdir(pop)

        if tess_ok and pop_ok:
            return True

        problems = []
        if not tess_ok:
            problems.append("• Tesseract (tesseract.exe) не задан или не найден")
        if not pop_ok:
            problems.append("• Poppler (папка bin) не задан или не найден")

        messagebox.showwarning(
            "Нужны настройки OCR",
            "Для распознавания PDF не хватает путей:\n\n"
            + "\n".join(problems)
            + "\n\nСейчас откроется окно настроек — укажите пути через «Обзор» и нажмите «Сохранить»."
        )

        self.open_settings()

        tess = (self.config.get("tesseract_path") or "").strip()
        pop = (self.config.get("poppler_path") or "").strip()
        tess_ok = bool(tess) and os.path.isfile(tess)
        pop_ok = bool(pop) and os.path.isdir(pop)

        return tess_ok and pop_ok

    # -----------------------------------------------------------------
    # Выбор и загрузка файла
    # -----------------------------------------------------------------
    def select_file(self):
        filetypes = [
            ("Все поддерживаемые", "*.xlsx *.xls *.docx *.doc *.pdf"),
            ("Excel файлы", "*.xlsx *.xls"),
            ("Word файлы", "*.docx *.doc"),
            ("PDF файлы", "*.pdf"),
            ("Все файлы", "*.*")
        ]

        filename = filedialog.askopenfilename(
            title="Выберите файл со списком участников",
            filetypes=filetypes,
            initialdir=os.path.dirname(os.path.abspath(__file__))
        )

        if filename:
            self.file_path = filename
            self.file_label.config(text=os.path.basename(filename), foreground="black")
            self.status_var.set(f"Загрузка данных из {os.path.basename(filename)}...")
            self.root.update()
            self.load_participants()

    def load_participants(self):
        try:
            ext = Path(self.file_path).suffix.lower()

            if ext in (".xlsx", ".xls"):
                excel_tables = extract_excel_tables(self.file_path)
                if len(excel_tables) == 1:
                    self.participants = excel_tables[0]["participants"]
                elif len(excel_tables) > 1:
                    chosen = self.choose_table_dialog(excel_tables)
                    self.participants = chosen if chosen is not None else []
                else:
                    self.participants = []
            elif ext == ".docx":
                self.participants = extract_from_docx(self.file_path)
            elif ext == ".doc":
                self.participants = extract_from_doc(self.file_path)
            elif ext == ".pdf":
                self.participants = self.extract_from_pdf_ocr()
            else:
                messagebox.showerror("Ошибка", f"Неподдерживаемый формат: {ext}")
                return

            self._normalize_participants_keys()

            if not self.participants:
                messagebox.showwarning("Внимание", "Не удалось извлечь участников из файла")
                self.status_var.set("Не удалось извлечь участников")
                return

            self.fill_table()

            self.status_var.set(f"Загружено участников: {len(self.participants)}")
            messagebox.showinfo("Готово", f"Загружено {len(self.participants)} участников")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка загрузки: {str(e)}")
            self.status_var.set("Ошибка загрузки")

    def _normalize_participants_keys(self):
        for p in self.participants:
            if "subdivision" not in p:
                p["subdivision"] = p.get("organization", "")
            p.pop("organization", None)

            p.setdefault("name", "")
            p.setdefault("position", "")
            p["group"] = self._normalize_group(p.get("group", ""))
            p.setdefault("permit_codes", "")
            p.setdefault("permit_text", "")

    def extract_from_pdf_ocr(self):
        try:
            from pdf_table_ocr import extract_tables_from_pdf
        except ImportError:
            messagebox.showerror(
                "Ошибка",
                "Не установлены зависимости для OCR:\n"
                "pip install pdf2image pytesseract opencv-python-headless numpy"
            )
            return []

        if not self._check_ocr_paths():
            self.status_var.set("Распознавание отменено: не заданы пути OCR")
            return []

        self.status_var.set("Распознавание таблиц в PDF (офлайн OCR)...")
        self.root.update()

        pages = self.pages_var.get().strip()

        try:
            cfg = get_config()
            if str(cfg.get("pdf_pages", "") or "").strip() != pages:
                cfg["pdf_pages"] = pages
                save_config(cfg)
        except Exception:
            pass

        try:
            tables = extract_tables_from_pdf(
                self.file_path,
                dpi=self.config.get("ocr_dpi"),
                lang=self.config.get("ocr_lang"),
                tesseract_path=self.config.get("tesseract_path"),
                poppler_path=self.config.get("poppler_path"),
                pages=pages,
            )
        except Exception as e:
            messagebox.showerror("Ошибка OCR", f"Ошибка распознавания: {str(e)}")
            return []

        if not tables:
            messagebox.showwarning(
                "Внимание",
                "Таблицы не найдены или не распознаны.\n"
                "Проверьте качество скана, язык и страницы в «Настройках OCR»."
            )
            return []

        if len(tables) == 1:
            return self._participants_from_table(tables[0])

        chosen = self.choose_table_dialog(tables)
        return chosen if chosen is not None else []

    def choose_table_dialog(self, tables):
        dialog = tk.Toplevel(self.root)
        dialog.title("Выберите таблицу")
        dialog.geometry("700x320")
        dialog.grab_set()

        tk.Label(
            dialog,
            text="В заявке найдено несколько таблиц. Выберите нужную:",
            font=("Arial", 11, "bold"),
            pady=10
        ).pack()

        listbox = tk.Listbox(dialog, font=("Arial", 10), height=8)

        for i, t in enumerate(tables, 1):
            group = self._normalize_group(t.get("group_guess", ""))
            hint = f" [группа: {group}]" if group else ""
            heading = (t["heading"][:60] + "...") if len(t["heading"]) > 60 else t["heading"]

            listbox.insert(
                tk.END,
                f"{i}. {heading or '(без заголовка)'}{hint} — участников: {len(t['participants'])}"
            )

        listbox.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        listbox.selection_set(0)

        result = {"participants": None}

        def on_ok():
            sel = listbox.curselection()
            if sel:
                result["participants"] = self._participants_from_table(tables[sel[0]])
            dialog.destroy()

        def on_all():
            combined = []
            for t in tables:
                combined.extend(self._participants_from_table(t))
            result["participants"] = combined
            dialog.destroy()

        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(pady=10)

        ttk.Button(
            btn_frame,
            text="Загрузить выбранную таблицу",
            command=on_ok
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            btn_frame,
            text="Загрузить все таблицы",
            command=on_all
        ).pack(side=tk.LEFT, padx=5)

        dialog.wait_window()

        return result["participants"]

    def _participants_from_table(self, table):
        heading_group = self._normalize_group(table.get("group_guess", ""))
        participants = []
        for p in table.get("participants", []):
            participant = dict(p)
            row_group = self._normalize_group(participant.get("group", ""))
            participant["group"] = row_group or heading_group
            participant.setdefault("permit_codes", "")
            participant.setdefault("permit_text", "")
            participants.append(participant)
        return participants

    # -----------------------------------------------------------------
    # Таблица: заполнение, редактирование, синхронизация
    # -----------------------------------------------------------------
    def fill_table(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

        for idx, p in enumerate(self.participants):
            self.tree.insert(
                "",
                tk.END,
                iid=str(idx),
                values=(
                    idx + 1,
                    p.get("name", ""),
                    p.get("position", ""),
                    p.get("subdivision", ""),
                    p.get("group", ""),
                    p.get("permit_codes", "")
                )
            )

    def add_participant(self):
        """Добавляет пустую строку участника в конец списка."""
        self._sync_all()
        self.participants.append({
            "name": "",
            "position": "",
            "subdivision": "",
            "group": "",
            "permit_codes": "",
            "permit_text": "",
            "permits_by_group": {},
            "permit_text_by_group": {},
        })
        self.fill_table()
        children = self.tree.get_children()
        if children:
            self.tree.selection_set(children[-1])
            self.tree.focus(children[-1])

    def remove_participants(self):
        """Удаляет выбранные строки участников."""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Внимание", "Сначала выдели строки, которые нужно удалить")
            return
        if not messagebox.askyesno(
            "Удаление",
            f"Удалить выбранных участников ({len(selected)})?"
        ):
            return
        self._sync_all()
        indices = sorted((int(i) for i in selected), reverse=True)
        for idx in indices:
            if 0 <= idx < len(self.participants):
                del self.participants[idx]
        self.fill_table()

    def _sync_participant(self, item):
        try:
            idx = int(item)
        except ValueError:
            return

        if idx < 0 or idx >= len(self.participants):
            return

        values = self.tree.item(item, "values")

        p = self.participants[idx]
        p["name"] = str(values[1]).strip()
        p["position"] = str(values[2]).strip()
        p["subdivision"] = str(values[3]).strip()
        p["group"] = self._normalize_group(values[4])
        p["permit_codes"] = str(values[5]).strip()

        groups_map = split_permit_groups(p["permit_codes"])
        p["permits_by_group"] = groups_map
        texts = {}
        for g, codes in groups_map.items():
            codes_list, _ = normalize_permit_codes(codes)
            texts[g] = "\n".join(PERMITS[c].strip() for c in codes_list)
        p["permit_text_by_group"] = texts
        p["permit_text"] = "\n".join(t for t in texts.values() if t)

    def _sync_all(self):
        for item in self.tree.get_children():
            self._sync_participant(item)

    def on_tree_double_click(self, event):
        region = self.tree.identify("region", event.x, event.y)
        if region != "cell":
            return

        item = self.tree.identify_row(event.y)
        column = self.tree.identify_column(event.x)

        if not item or not column:
            return

        col_index = int(column.replace("#", "")) - 1

        if col_index == 0:
            return

        field = self.tree["columns"][col_index]
        current_value = self.tree.set(item, field)

        field_labels = {
            "name": "ФИО",
            "position": "Должность",
            "subdivision": "Подразделение",
            "group": "Группа",
            "permits": "Коды допусков"
        }

        label = field_labels.get(field, field)

        if field == "permits":
            self.edit_permits_item(item)
            return

        if field == "group":
            new_value = askstring(
                "Группа",
                "Введите группу:\n1, 2, 3 или несколько: 2,3 / 1 и 2",
                initialvalue=current_value,
                parent=self.root
            )
        else:
            new_value = askstring(
                f"Редактирование: {label}",
                f"Исправьте значение поля «{label}»:",
                initialvalue=current_value,
                parent=self.root
            )

        if new_value is None:
            return

        if field == "group":
            new_value = self._normalize_group(new_value)
        else:
            new_value = new_value.strip()

        self.tree.set(item, field, new_value)
        self._sync_participant(item)

    def edit_permits_item(self, item):
        current_codes = self.tree.set(item, "permits")
        name = self.tree.set(item, "name")

        try:
            person_groups = self._split_groups(self.participants[int(item)].get("group", ""))
        except (ValueError, IndexError):
            person_groups = []

        dialog = PermitEditorDialog(
            self.root,
            current_codes=current_codes,
            title=f"Допуски: {name}",
            default_groups=person_groups or None
        )

        if dialog.result is None:
            return

        self.tree.set(item, "permits", dialog.result)
        self._sync_participant(item)

    def on_tree_right_click(self, event):
        item = self.tree.identify_row(event.y)
        selection = self.tree.selection()

        if item and item not in selection:
            self.tree.selection_set(item)

        if not self.tree.selection():
            return

        menu = tk.Menu(self.root, tearoff=0)
        menu.add_command(
            label="Назначить допуски выбранным",
            command=self.mass_assign_permits
        )
        menu.add_command(
            label="Назначить группу выбранным",
            command=self.mass_assign_group
        )
        menu.add_command(
            label="Очистить допуски у выбранных",
            command=self.clear_selected_permits
        )
        menu.tk_popup(event.x_root, event.y_root)

    def mass_assign_permits(self):
        selected = self.tree.selection()
        if not selected:
            return

        first_item = selected[0]
        current_codes = self.tree.set(first_item, "permits")

        dialog = PermitEditorDialog(
            self.root,
            current_codes=current_codes,
            title="Допуски для выбранных участников"
        )

        if dialog.result is None:
            return

        for item in selected:
            self.tree.set(item, "permits", dialog.result)
            self._sync_participant(item)

    def mass_assign_group(self):
        selected = self.tree.selection()
        if not selected:
            return

        first_item = selected[0]
        current_group = self.tree.set(first_item, "group")

        new_group = askstring(
            "Группа для выбранных участников",
            "Введите группу:\n1, 2, 3 или несколько: 2,3 / 1 и 2",
            initialvalue=current_group,
            parent=self.root
        )

        if new_group is None:
            return

        new_group = self._normalize_group(new_group)

        for item in selected:
            self.tree.set(item, "group", new_group)
            self._sync_participant(item)

    def clear_selected_permits(self):
        for item in self.tree.selection():
            self.tree.set(item, "permits", "")
            self._sync_participant(item)

    # -----------------------------------------------------------------
    # Вспомогательные функции по группам и именам файлов
    # -----------------------------------------------------------------
    def _normalize_group(self, group):
        g = str(group or "").strip().lower()
        if not g:
            return ""

        g = g.replace("ё", "е")
        g = re.sub(r"(\d)\s*\.\s*(\d)", r"\1,\2", g)
        gc = re.sub(r"\s+", "", g)

        if gc in ("1", "первая", "1группа", "1группы", "гр.1", "1гр", "гр1"):
            return "1"
        if gc in ("2", "вторая", "2группа", "2группы", "гр.2", "2гр", "гр2"):
            return "2"
        if gc in ("3", "третья", "3группа", "3группы", "гр.3", "3гр", "гр3"):
            return "3"
        if gc in ("1и2", "1,2", "1-2", "1и2группа", "1и2группы", "1,2группы"):
            return "1 и 2"

        return g

    def _split_groups(self, group_str):
        """Разбивает '2,3', '1 и 2', '1и2', '2 3' на отдельные группы: ['2','3']."""
        s = str(group_str or "").strip().lower().replace("ё", "е")
        if not s:
            return []
        s = re.sub(r"\s+", "", s)
        parts = re.split(r"[,;./]|и", s)
        groups = []
        for part in parts:
            if part in ("1", "2", "3") and part not in groups:
                groups.append(part)
        return groups

    def _course_group_phrase(self, group):
        g = self._normalize_group(group)
        if not g:
            return "группы"
        if g in ("1 и 2", "1и2"):
            return "1 и 2 групп"
        return f"{g} группы"

    def _sanitize_filename(self, name):
        return re.sub(r'[\\/:*?"<>|]', "_", str(name)).strip()

    # -----------------------------------------------------------------
    # Справка по допускам
    # -----------------------------------------------------------------
    def show_permits_help(self):
        help_window = tk.Toplevel(self.root)
        help_window.title("Справка по допускам")
        help_window.geometry("650x520")

        tk.Label(
            help_window,
            text="Коды допусков по листам регистрации",
            font=("Arial", 12, "bold"),
            pady=10
        ).pack()

        text_frame = ttk.Frame(help_window)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        text_widget = tk.Text(text_frame, wrap=tk.WORD, font=("Consolas", 10))
        scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)

        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        for key in ("1", "2", "3"):
            text_widget.insert(tk.END, f"=== ГРУППА {key} ===\n\n")
            for code in SHEET_GROUPS[key]:
                text_widget.insert(tk.END, f"  {code:2d} - {str(PERMITS[code]).strip()}\n")
            text_widget.insert(tk.END, "\n")

        text_widget.insert(tk.END, "=== ФОРМАТ ВВОДА ===\n\n")
        text_widget.insert(tk.END, "  • 2: 8,10,11 - допуски группы 2\n")
        text_widget.insert(tk.END, "  • 2: 8,10; 3: 5,6 - несколько групп\n")

        text_widget.config(state=tk.DISABLED)

        ttk.Button(help_window, text="Закрыть", command=help_window.destroy).pack(pady=10)

    # -----------------------------------------------------------------
    # Генерация протокола (одного или нескольких по группам)
    # -----------------------------------------------------------------
    def generate_protocol(self):
        if not self.participants:
            messagebox.showwarning("Внимание", "Сначала загрузите файл со списком участников")
            return

        if not self.protocol_num.get().strip():
            messagebox.showwarning("Внимание", "Введите номер протокола")
            return

        self._sync_all()

        missing_permits = []
        for p in self.participants:
            if not str(p.get("permit_codes", "")).strip():
                missing_permits.append(p.get("name", ""))

        if missing_permits:
            result = messagebox.askyesno(
                "Внимание",
                "Не введены допуски для:\n"
                + "\n".join(missing_permits[:5])
                + ("\n..." if len(missing_permits) > 5 else "")
                + "\n\nПродолжить без допусков?"
            )
            if not result:
                return

        base_protocol_info = {
            "protocol_number": self.protocol_num.get().strip(),
            "date": self.date_entry.get().strip() or datetime.now().strftime("%d.%m.%Y"),
            "course_name": "Безопасные методы и приемы выполнения работ на высоте",
            "group": self.group_var.get(),
            "hours": self.hours_var.get(),
            "organization": self.org_var.get().strip(),
        }

        groups = []
        for p in self.participants:
            person_groups = self._split_groups(p.get("group", ""))
            if not person_groups:
                person_groups = [self._normalize_group(base_protocol_info["group"])]
            for g in person_groups:
                if g not in groups:
                    groups.append(g)

        if len(groups) > 1:
            answer = messagebox.askyesno(
                "Найдено несколько групп",
                "Участники относятся к группам:\n\n"
                + ", ".join(groups)
                + "\n\nСделать ОТДЕЛЬНЫЕ протоколы по группам?\n"
                + "«Да» — как для Транснефти. «Нет» — один общий протокол, "
                + "и человек с несколькими группами получит по строке на группу."
            )
            if answer:
                self._generate_separate_protocols(groups, base_protocol_info)
                return

        self.protocol_info = dict(base_protocol_info)

        if len(groups) == 1:
            self.protocol_info["group"] = groups[0]

        group_override = None

        safe_name = self._sanitize_filename(
            f"Протокол_{self.protocol_info['protocol_number']}_"
            f"{self.protocol_info['date']}_"
            f"группа_{self.protocol_info['group']}.docx"
        )
        if not safe_name.lower().endswith(".docx"):
            safe_name += ".docx"

        output_path = filedialog.asksaveasfilename(
            title="Сохранить протокол",
            defaultextension=".docx",
            filetypes=[("Word документ", "*.docx")],
            initialdir=os.path.dirname(os.path.abspath(__file__)),
            initialfile=safe_name
        )

        if not output_path:
            return

        try:
            self.status_var.set("Создание протокола...")
            self.root.update()

            self.create_protocol_file(
                output_path,
                protocol_info=self.protocol_info,
                group_override=group_override
            )

            self.status_var.set(f"Протокол сохранён: {os.path.basename(output_path)}")
            messagebox.showinfo(
                "Готово",
                f"Протокол успешно создан!\n\n{output_path}"
            )
        except Exception as e:
            messagebox.showerror(
                "Ошибка",
                f"Ошибка создания протокола: {str(e)}"
            )
            self.status_var.set("Ошибка создания протокола")

    def _generate_separate_protocols(self, groups, base_protocol_info):
        output_dir = filedialog.askdirectory(
            title="Выберите папку для сохранения протоколов по группам"
        )
        if not output_dir:
            return

        created_files = []

        try:
            for group in groups:
                group_participants = [
                    p for p in self.participants
                    if group in (
                        self._split_groups(p.get("group", ""))
                        or [self._normalize_group(base_protocol_info["group"])]
                    )
                ]
                if not group_participants:
                    continue

                protocol_info = dict(base_protocol_info)
                protocol_info["group"] = group

                safe_group = self._sanitize_filename(str(group).replace(" ", "_"))

                filename = self._sanitize_filename(
                    f"Протокол_{protocol_info['protocol_number']}_"
                    f"{protocol_info['date']}_"
                    f"группа_{safe_group}.docx"
                )
                if not filename.lower().endswith(".docx"):
                    filename += ".docx"

                output_path = os.path.join(output_dir, filename)

                self.status_var.set(f"Создание протокола для группы {group}...")
                self.root.update()

                self.create_protocol_file(
                    output_path,
                    participants=group_participants,
                    protocol_info=protocol_info,
                    group_override=group
                )

                created_files.append(output_path)

            if created_files:
                self.status_var.set(f"Создано протоколов: {len(created_files)}")
                messagebox.showinfo(
                    "Готово",
                    "Созданы протоколы:\n\n" + "\n".join(created_files)
                )
            else:
                messagebox.showwarning("Внимание", "Не удалось создать протоколы по группам")

        except Exception as e:
            messagebox.showerror(
                "Ошибка",
                f"Ошибка создания протоколов: {str(e)}"
            )
            self.status_var.set("Ошибка создания протоколов")

    # -----------------------------------------------------------------
    # Создание Word-файла протокола (формат как в образце 19.06.2026)
    # -----------------------------------------------------------------
    def _date_parts(self, date_str):
        """Разбирает дату "ДД.ММ.ГГГГ" и возвращает (день, месяц словами, год) или None."""
        months = [
            "января", "февраля", "марта", "апреля", "мая", "июня",
            "июля", "августа", "сентября", "октября", "ноября", "декабря",
        ]
        s = str(date_str or "").strip()
        m = re.match(r"^(\d{1,2})\.(\d{1,2})\.(\d{4})$", s)
        if not m:
            return None
        day, month, year = m.groups()
        if not 1 <= int(month) <= 12:
            return None
        return int(day), months[int(month) - 1], year

    def create_protocol_file(
        self,
        output_path,
        participants=None,
        protocol_info=None,
        group_override=None
    ):
        import docx
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        if participants is None:
            participants = self.participants
        if protocol_info is None:
            protocol_info = self.protocol_info

        org_name = (protocol_info.get("organization") or "").strip()

        # Дата словами: "19.06.2026" -> "от «19» июня 2026 г."
        parts = self._date_parts(protocol_info["date"])
        if parts:
            day, month_name, year = parts
            date_line = f"от «{day}» {month_name} {year} г."
            vedomost_date = f'от "{day}" {month_name} {year} г.'
        else:
            date_line = f'от «{protocol_info["date"]}» г.'
            vedomost_date = f'от "{protocol_info["date"]}" г.'

        doc = docx.Document()

        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.0)

        # === Шапка ===
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('ЧУДПО "Учебный комбинат "Мелиоратор"')
        run.font.size = Pt(14)
        run.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'П Р О Т О К О Л    № {protocol_info["protocol_number"]}')
        run.font.size = Pt(14)
        run.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(date_line)
        run.font.size = Pt(12)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("заседания экзаменационной комиссии")
        run.font.size = Pt(12)

        # === Состав комиссии: все строки с начала строки ===
        doc.add_paragraph("в составе:")
        doc.add_paragraph("председатель комиссии: директор ЧУДПО УК «Мелиоратор» - Первушов Александр Алексеевич")
        doc.add_paragraph("члены комиссии: заместитель директора ЧУДПО УК «Мелиоратор» - Украинцев Павел Алексеевич")
        doc.add_paragraph("специалист по охране труда ЧУДПО УК «Мелиоратор» – Леонтьева Лариса Петровна")

        # === Описание курса (как в образце) ===
        course_desc = (
            f'по выпуску окончивших обучение по курсу: «{protocol_info["course_name"]}» '
            f'в соответствии с Правилами по охране труда при работе на высоте '
            f'(приказ Минтруда России от 16.11.2020г. № 782н) '
            f'Произведя проверку знаний обучившихся по программе в объеме {protocol_info["hours"]} часов '
            f'комиссия постановила:'
        )
        doc.add_paragraph(course_desc)

        doc.add_paragraph("")

        # === Основная таблица ===
        rows_plan = []
        for participant in participants:
            if group_override:
                sub_groups = [group_override]
            else:
                sub_groups = self._split_groups(participant.get("group", ""))
                if not sub_groups:
                    sub_groups = [protocol_info.get("group", "")]
            rows_plan.append((participant, sub_groups))

        total_rows = sum(len(sg) for _, sg in rows_plan)

        table = doc.add_table(rows=1 + total_rows, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Ширины столбцов основной таблицы - примерно как в образце:
        # узкие служебные, широкий "Допущен к работам" (в сумме 17 см).
        # Порядок: №, ФИО+организация, должность, заключение,
        #          № удостоверения, группа, допуски.
        table.autofit = False
        table.allow_autofit = False
        widths = [Cm(1.1), Cm(3.0), Cm(3.5), Cm(1.7), Cm(2.0), Cm(2.0), Cm(5.2)]
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w

        headers = [
            "№\nп/п",
            "фамилия, имя, отчество,\nорганизация",
            "Профессия\nдолжность",
            "заключение комиссии\nсдано/\nне сдано",
            "№ выданного\nудостоверения",
            "Группа\nпо безопасности\nработ на высоте",
            "Допущен к работам:"
        ]

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.bold = True

        row_idx = 1
        for idx, (participant, sub_groups) in enumerate(rows_plan, 1):
            first_row = row_idx

            for gi, grp in enumerate(sub_groups):
                row = table.rows[row_idx]

                if gi == 0:
                    row.cells[0].text = f"{idx}."

                    # Под ФИО - подразделение, а если его нет - организация.
                    # Под должностью ничего не пишется.
                    #name_org = participant.get("name", "")
                    second_line = (participant.get("subdivision") or "").strip() or org_name
                    if second_line:
                        name_org += "\n" + second_line
                    row.cells[1].text = name_org

                    row.cells[2].text = participant.get("position", "")

                row.cells[3].text = "сдано"
                row.cells[4].text = "____"
                row.cells[5].text = grp

                texts = participant.get("permit_text_by_group", {}) or {}
                full_permits = texts.get(grp) or texts.get("") or ""
                if not full_permits:
                    full_permits = participant.get("permit_text", "")
                row.cells[6].text = full_permits

                row_idx += 1

            # Несколько групп - объединяем №, ФИО и должность в общие ячейки.
            if len(sub_groups) > 1:
                for j in (0, 1, 2):
                    top_cell = table.cell(first_row, j)
                    for r in range(first_row + 1, row_idx):
                        top_cell = top_cell.merge(table.cell(r, j))

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # === Подписи после таблицы ===
        doc.add_paragraph("")
        for line in (
            "Председатель комиссии: _______________ А.А. Первушов",
            "Члены комиссии: ________________ П.А. Украинцев",
            "_______________ Л.П. Леонтьева",
        ):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(line)

        # === Ведомость после протокола (люди расписываются) ===
        doc.add_paragraph("")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'В Е Д О М О С Т Ь к протоколу № {protocol_info["protocol_number"]}')
        run.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('«Безопасность работ на высоте»')
        run.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Конец обучения    {vedomost_date}')

        doc.add_paragraph("")

        v_table = doc.add_table(rows=1 + len(participants), cols=4)
        v_table.style = 'Table Grid'
        v_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Узкий столбец №, широкий столбец ФИО (в сумме 17 см - ширина поля).
        v_table.autofit = False
        v_table.allow_autofit = False
        v_widths = [Cm(1.5), Cm(9.5), Cm(4.3), Cm(3.2)]
        for row in v_table.rows:
            for i, w in enumerate(v_widths):
                row.cells[i].width = w

        v_headers = [
            "№\nп/п",
            "Фамилия, имя, отчество",
            "№ свидетельства\n/удостоверения/",
            "Подпись обучающегося"
        ]

        for i, header in enumerate(v_headers):
            cell = v_table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.bold = True

        for idx, participant in enumerate(participants, 1):
            row = v_table.rows[idx]
            row.cells[0].text = f"{idx}."

            name_org = participant.get("name", "")
            second_line = (participant.get("subdivision") or "").strip() or org_name
            if second_line:
                name_org += "\n" + second_line
            row.cells[1].text = name_org

            row.cells[2].text = ""
            row.cells[3].text = ""

        for row in v_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        # Межстрочный интервал 1 для всего документа.
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.0
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing = 1.0

        doc.save(output_path)

    # -----------------------------------------------------------------
    # Очистка
    # -----------------------------------------------------------------
    def clear_all(self):
        self.file_path = None
        self.participants = []
        self.permits_data = {}
        self.protocol_info = {}

        self.file_label.config(text="Файл не выбран", foreground="gray")

        self.protocol_num.delete(0, tk.END)

        self.date_entry.delete(0, tk.END)
        self.date_entry.insert(0, datetime.now().strftime("%d.%m.%Y"))

        self.group_var.set("3")
        self.hours_var.set("32")
        self.org_var.set("")

        for item in self.tree.get_children():
            self.tree.delete(item)

        self.status_var.set("Готово к работе")


def main():
    root = tk.Tk()
    app = ProtocolBotApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()