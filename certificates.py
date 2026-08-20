"""
certificates.py - удостоверения ПО ШАБЛОНУ Word с плейсхолдерами.

Вёрстка живёт в файле certificate_template.docx рядом с программой и
НЕ ТРОГАЕТСЯ программой: все пустые строки и отступы остаются как в шаблоне.
Программа только:
  1) подставляет данные вместо {плейсхолдеров};
  2) если текст допусков длиннее, чем строка в шаблоне, - добирает строки
     из пустого резерва МЕЖДУ допусками и подписью (подпись остаётся внизу);
  3) фиксирует ширины колонок и уменьшает шрифт очень длинной должности.

Плейсхолдеры: {number} {surname} {name} {patronymic} {position} {org}
{date_issue} {date_valid} {group} {permit_text} {protocol_number}
{protocol_date} {hours}

2 удостоверения на страницу. Группы 1 и 2 - на 3 года, группа 3 - на 5 лет.
Часы практического обучения - всегда 5.
"""
import copy
import math
import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TEMPLATE_NAME = "certificate_template.docx"
PRACTICE_HOURS = "5"

# Геометрия таблицы (как в эталоне) - фиксируем, чтобы длинный
# текст не расползался по ширине.
LEFT_WIDTH_CM = 8.93
RIGHT_WIDTH_CM = 8.75
# Граница слева для должности/организации: левее - место под фото.
PHOTO_SPACE_CM = 3.5

# Примерная вместимость одной строки правой ячейки (для расчёта допусков).
CHARS_PER_LINE = 65
# Примерная вместимость одной строки для должности в левой ячейке.
POSITION_CHARS_PER_LINE = 45

SIGNATURE_MARKERS = ("Руководитель", "Первушов", "М.П.", "(подпись)")
PERMIT_MARKER = "Может быть допущен"
BASE_MARKER = "Основание"

_MONTHS = [
    "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]

# Схемный порядок элементов - чтобы Word не считал файл битым.
_TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
    "tblStyleRowBandSize", "tblStyleColBandSize", "tblW", "jc",
    "tblCellSpacing", "tblInd", "tblBorders", "shd", "tblLayout",
    "tblCellMar", "tblLook", "caption", "tblDescription",
]
_TCPR_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders",
    "shd", "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign",
    "hideMark",
]


def _tw(cm):
    return str(int(round(cm * 567)))


def _name(tag):
    return tag.split("}")[-1]


def _date_parts(date_str):
    m = re.match(r"^(\d{1,2})\.(\d{1,2})\.(\d{4})$", str(date_str or "").strip())
    if not m:
        return None
    day, month, year = m.groups()
    if not 1 <= int(month) <= 12:
        return None
    return int(day), _MONTHS[int(month) - 1], int(year)


def _next_number(current):
    m = re.match(r"^(\d+)(.*)$", str(current).strip(), re.S)
    if not m:
        return current
    return str(int(m.group(1)) + 1) + m.group(2)


def _upsert(parent, tag, order):
    el = parent.find(qn(tag))
    if el is not None:
        return el
    el = OxmlElement(tag)
    pos = order.index(_name(tag)) if _name(tag) in order else 999
    for i, child in enumerate(parent):
        cn = _name(child.tag)
        if cn in order and order.index(cn) > pos:
            parent.insert(i, el)
            return el
    parent.append(el)
    return el


def _strip_ids(element, seed=0):
    """Убирает дублирующиеся у копий paraId/textId и делает уникальными
    id фигур и docPr - иначе Word считает файл повреждённым."""
    n = 1000 + seed * 50
    for el in element.iter():
        for k in list(el.attrib):
            if k.split("}")[-1] in ("paraId", "textId"):
                del el.attrib[k]
        ln = _name(el.tag)
        if ln == "docPr":
            el.set("id", str(n))
            n += 1
        elif ln == "shape":
            el.set("id", f"_x0000_i{n}")
            n += 1


def _merge_split_placeholders(p):
    """Склеивает раны, в которых Word разбил плейсхолдер на куски."""
    runs = [r for r in p.findall(qn("w:r")) if r.find(qn("w:t")) is not None]
    i = 0
    while i < len(runs):
        t = runs[i].find(qn("w:t"))
        txt = t.text or ""
        j = i
        while txt.count("{") > txt.count("}") and j + 1 < len(runs):
            j += 1
            txt += (runs[j].find(qn("w:t")).text or "")
        if j > i:
            t.text = txt
            t.set(qn("xml:space"), "preserve")
            for k in range(i + 1, j + 1):
                p.remove(runs[k])
            runs = [r for r in p.findall(qn("w:r")) if r.find(qn("w:t")) is not None]
        i += 1


def _fill_element(element, mapping):
    """Заменяет {плейсхолдеры} во всех текстовых узлах таблицы."""
    for p in element.iter(qn("w:p")):
        _merge_split_placeholders(p)
    for t in element.iter(qn("w:t")):
        text = t.text
        if not text or "{" not in text:
            continue
        for key, value in mapping.items():
            text = text.replace("{" + key + "}", value)
        t.text = text


def _fix_geometry(table):
    """Фиксированная раскладка + точные ширины колонок (схемобезопасно)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    lay = _upsert(tblPr, "w:tblLayout", _TBLPR_ORDER)
    lay.set(qn("w:type"), "fixed")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        cols = grid.findall(qn("w:gridCol"))
        if len(cols) == 2:
            cols[0].set(qn("w:w"), _tw(LEFT_WIDTH_CM))
            cols[1].set(qn("w:w"), _tw(RIGHT_WIDTH_CM))
    for row in table.rows:
        for idx, w in enumerate((LEFT_WIDTH_CM, RIGHT_WIDTH_CM)):
            if idx >= len(row.cells):
                continue
            tcPr = row.cells[idx]._tc.get_or_add_tcPr()
            tcW = _upsert(tcPr, "w:tcW", _TCPR_ORDER)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), _tw(w))


def _walk_cells(cells):
    for cell in cells:
        for p in cell.paragraphs:
            yield p
        for nt in cell.tables:
            for nr in nt.rows:
                for p in _walk_cells(nr.cells):
                    yield p


_PPR_ORDER = [
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd",
    "tabs", "suppressAutoHyphens", "kinsoku", "wordWrap",
    "overflowPunct", "topLinePunct", "autoSpaceDE", "autoSpaceDN",
    "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind",
    "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc",
    "textDirection", "textAlignment", "textboxProperties", "outlineLvl",
    "divId", "cnfStyle",
]


def _suppress_auto_hyphens(p):
    pEl = p._p                      # внутренний XML-элемент абзаца
    pPr = pEl.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        pEl.insert(0, pPr)
    _upsert(pPr, "w:suppressAutoHyphens", _PPR_ORDER)

def _shrink_long_position(table, position):
    """Длинная должность - шрифт мельче, запрет переносов посреди слова
    и граница слева: не залезает на место под фото."""
    if not position:
        return
    size = 8 if len(position) <= 40 else (7 if len(position) <= 55 else 6)
    for row in table.rows:
        for p in _walk_cells(row.cells):
            for r in p.runs:
                if position in (r.text or ""):
                    r.font.size = Pt(size)
                    _suppress_auto_hyphens(p)
                    p.paragraph_format.left_indent = Cm(PHOTO_SPACE_CM)

def _consume_spacers_for_permit(table, mapping):
    """ЕДИНСТВЕННАЯ правка макета: если текст допусков занимает больше
    строк, чем в шаблоне, недостающее место берётся из пустых строк,
    стоящих МЕЖДУ допусками и подписью. Остальной шаблон не трогается."""
    if not table.rows or len(table.rows[0].cells) < 2:
        return
    tc = table.rows[0].cells[1]._tc

    permit = mapping.get("permit_text", "") or ""
    intro = "Может быть допущен (а) к работе на высоте: "
    extra = max(0, math.ceil((len(intro) + len(permit)) / CHARS_PER_LINE) - 1)
    if not extra:
        return

    base_idx = permit_idx = end = None
    for i, child in enumerate(tc):
        txt = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if permit_idx is None and PERMIT_MARKER in txt:
            permit_idx = i
        if base_idx is None and BASE_MARKER in txt:
            base_idx = i
        if any(m in txt for m in SIGNATURE_MARKERS):
            end = i
            break
    # резерв едим только между "Основанием:" и подписью
    start = base_idx if base_idx is not None else permit_idx
    if start is None or end is None or end <= start + 1:
        return

    spacers = []
    for i in range(start + 1, end):
        child = tc[i]
        if child.tag == qn("w:p"):
            text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
            if not text:
                spacers.append(child)

    for p_el in spacers[:extra]:
        p_el.getparent().remove(p_el)

def _consume_spacers_for_position(table, position):
    """Если должность переносится на доп. строки - съедает столько же
    пустых строк резерва в левой ячейке (с конца, ближе к таблице дат),
    чтобы высота удостоверения не менялась."""
    if not position or not table.rows or len(table.rows[0].cells) < 2:
        return
    extra = max(0, math.ceil(len(position) / POSITION_CHARS_PER_LINE) - 1)
    if not extra:
        return
    tc = table.rows[0].cells[0]._tc
    start = None
    for i, child in enumerate(tc):
        txt = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if position in txt:
            start = i
            break
    if start is None:
        return
    spacers = []
    for i in range(start + 1, len(tc)):
        child = tc[i]
        if child.tag == qn("w:p"):
            text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
            if not text:
                spacers.append(child)
    for p_el in spacers[-extra:]:
        p_el.getparent().remove(p_el)

def _ensure_trailing_p(tc):
    """Ячейка обязана заканчиваться абзацем, иначе Word считает файл битым."""
    if len(tc) == 0 or tc[-1].tag != qn("w:p"):
        tc.append(OxmlElement("w:p"))


def _append_before_sectpr(doc, element):
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    if sect is not None:
        sect.addprevious(element)
    else:
        body.append(element)


def create_certificates_file(output_path, participants_with_groups, info):
    template_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)), TEMPLATE_NAME
    )
    if not os.path.isfile(template_path):
        raise FileNotFoundError(
            f"Рядом с программой нет файла {TEMPLATE_NAME}.\n"
            "Создай шаблон с плейсхолдерами: одна таблица = одно удостоверение."
        )

    doc = Document(template_path)
    if not doc.tables:
        raise ValueError(f"В шаблоне {TEMPLATE_NAME} не найдена таблица удостоверения.")

    # Запоминаем "чистую" таблицу шаблона и снимаем обтекание "вокруг" -
    # иначе копии удостоверений накладываются друг на друга.
    clean_tbl = copy.deepcopy(doc.tables[0]._tbl)
    _tblPr = clean_tbl.find(qn("w:tblPr"))
    if _tblPr is not None:
        _ppr = _tblPr.find(qn("w:tblPPr"))
        if _ppr is not None:
            _tblPr.remove(_ppr)

    # Полностью очищаем тело документа (кроме свойств страницы),
    # чтобы не оставалось лишних абзацев и пустых половин листа.
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    parts = _date_parts(info.get("date"))
    if parts:
        day, month_name, year = parts
    else:
        now = datetime.now()
        day, month_name, year = now.day, _MONTHS[now.month - 1], now.year
    issue_date = f"«{day}» {month_name} {year} г."

    number = str(info.get("start_number", "")).strip()
    total = sum(len(groups) for _, groups in participants_with_groups)
    done = 0

    for participant, groups in participants_with_groups:
        words = (participant.get("name") or "").split()
        surname = words[0] if words else ""
        first_name = words[1] if len(words) > 1 else ""
        patronymic = " ".join(words[2:]) if len(words) > 2 else ""
        position = (participant.get("position") or "").strip()
        org = (participant.get("subdivision") or "").strip() or info.get("organization", "")
        texts = participant.get("permit_text_by_group", {}) or {}

        for grp in groups:
            try:
                valid_year = int(year) + (5 if grp == "3" else 3)
            except (TypeError, ValueError):
                valid_year = year

            mapping = {
                "number": str(number),
                "surname": surname,
                "name": first_name,
                "patronymic": patronymic,
                "position": position,
                "org": org,
                "date_issue": issue_date,
                "date_valid": f"«{day}» {month_name} {valid_year} г.",
                "group": grp,
                "permit_text": texts.get(grp) or texts.get("") or participant.get("permit_text", ""),
                "protocol_number": str(info.get("protocol_number", "")),
                "protocol_date": issue_date,
                "hours": PRACTICE_HOURS,
            }

            tbl = copy.deepcopy(clean_tbl)
            _strip_ids(tbl, done)
            _fill_element(tbl, mapping)

            table = Table(tbl, doc)
            _fix_geometry(table)
            _shrink_long_position(table, position)
            _consume_spacers_for_position(table, position)
            _consume_spacers_for_permit(table, mapping)
            _ensure_trailing_p(table.rows[0].cells[0]._tc)
            _ensure_trailing_p(table.rows[0].cells[1]._tc)

            _append_before_sectpr(doc, tbl)

            done += 1
            if done < total:
                if done % 2 == 0:
                    doc.add_page_break()          # 2 удостоверения на страницу
                else:
                    spacer = doc.add_paragraph("")
                    spacer.paragraph_format.space_after = Pt(6)
        number = _next_number(number)

    doc.save(output_path)