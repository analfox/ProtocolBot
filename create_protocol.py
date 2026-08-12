"""
ProtocolBot - Автоматизация создания протоколов об обучении

Использование:
    python create_protocol.py <файл_заявки> [номер_протокола] [дата] [группа] [часы]

Поддерживаемые форматы входных файлов:
    - Excel (.xlsx, .xls) - список участников
    - Word (.docx) - список участников
    - PDF - потребуется ручной ввод (или OCR с Tesseract)

Примеры:
    python create_protocol.py список.xlsx 543 24.07.2026 3 32
    python create_protocol.py заявка.docx
"""

import sys
import os
import re
from datetime import datetime
from pathlib import Path

from permits import PERMITS, parse_permits, list_permits


# ============================================================
# Извлечение данных из разных форматов
# ============================================================

def extract_from_excel(file_path: str) -> list:
    """Извлекает список участников из Excel-файла.
    .xlsx читает openpyxl, старые .xls - xlrd (их часто шлёт РЖД)."""
    ext = Path(file_path).suffix.lower()
    if ext == ".xls":
        return _extract_from_xls(file_path)

    import openpyxl

    participants = []
    wb = openpyxl.load_workbook(file_path, data_only=True)

    for sheet in wb.sheetnames:
        ws = wb[sheet]

        # Ищем заголовки
        header_row = None
        name_col = None
        position_col = None
        org_col = None

        for row_idx, row in enumerate(ws.iter_rows(max_row=10, values_only=True), 1):
            for col_idx, cell in enumerate(row):
                if cell and isinstance(cell, str):
                    cell_lower = cell.lower()
                    if "фамилия" in cell_lower or "фио" in cell_lower or "имя" in cell_lower:
                        header_row = row_idx
                        name_col = col_idx
                    elif "должность" in cell_lower or "профессия" in cell_lower:
                        position_col = col_idx
                    elif "организация" in cell_lower or "наименование" in cell_lower or "подразделение" in cell_lower:
                        org_col = col_idx
            if header_row:
                break

        if header_row is None or name_col is None:
            continue

        # Извлекаем данные
        for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
            if not row or not row[name_col]:
                continue

            name = str(row[name_col]).strip()
            if not name or len(name) < 3:
                continue

            position = ""
            if position_col is not None and len(row) > position_col and row[position_col]:
                position = str(row[position_col]).strip()

            org = ""
            if org_col is not None and len(row) > org_col and row[org_col]:
                org = str(row[org_col]).strip()

            participants.append({
                "name": name,
                "position": position,
                "organization": org,
            })

    return participants

def _norm_header(text) -> str:
    """Приводит заголовок к нижнему регистру без точек/пробелов/скобок:
    'Ф.И.О. (полностью)' -> 'фиополностью', чтобы 'фио' находилось."""
    return re.sub(r"[^а-яёa-z0-9]", "", str(text or "").lower())


def _cell_text(value) -> str:
    """Текст ячейки; числа вида 123.0 превращает в 123."""
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()

def _looks_like_person(name: str) -> bool:
    """Похоже ли на ФИО человека: 2-4 слова, из них >=2 с заглавной буквы.
    Так отсеиваются строки второй шапки: "СПИСОК", "профессионального
    развития...", даты, "Ф.И.О. (полностью)" и т.п."""
    words = (name or "").split()
    if not (2 <= len(words) <= 4):
        return False
    cap = sum(1 for w in words if w and (w[0].isupper() or w[0] == "Ё"))
    return cap >= 2

def extract_from_excel(file_path: str) -> list:
    """Извлекает список участников из Excel-файла.
    .xlsx читает openpyxl, старые .xls - xlrd (их часто шлёт РЖД)."""
    ext = Path(file_path).suffix.lower()
    if ext == ".xls":
        return _extract_from_xls(file_path)

    import openpyxl

    participants = []
    wb = openpyxl.load_workbook(file_path, data_only=True)

    for sheet in wb.sheetnames:
        ws = wb[sheet]

        header_row = None
        name_col = None
        position_col = None
        org_col = None
        group_col = None

        for row_idx, row in enumerate(ws.iter_rows(max_row=10, values_only=True), 1):
            for col_idx, cell in enumerate(row):
                hn = _norm_header(cell)
                if not hn:
                    continue
                if "фамилия" in hn or "фио" in hn or "имя" in hn:
                    header_row = row_idx
                    name_col = col_idx
                elif "должность" in hn or "профессия" in hn:
                    position_col = col_idx
                elif "подразделение" in hn or "организация" in hn or "наименование" in hn or hn == "сп":
                    org_col = col_idx
                elif "группа" in hn:
                    group_col = col_idx
            if header_row:
                break

        if header_row is None or name_col is None:
            continue

        for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
            if not row or name_col >= len(row):
                continue

            name = _cell_text(row[name_col])
            if not name or len(name) < 3:
                continue
            # Пропускаем мусор (вторая шапка и т.п.) - берём только похожее на ФИО.
            if not _looks_like_person(name):
                continue

            position = _cell_text(row[position_col]) if position_col is not None and len(row) > position_col else ""
            org = _cell_text(row[org_col]) if org_col is not None and len(row) > org_col else ""
            group = _cell_text(row[group_col]) if group_col is not None and len(row) > group_col else ""

            participants.append({
                "name": name,
                "position": position,
                "organization": org,
                "group": group,
            })

    return participants


def _extract_from_xls(file_path: str) -> list:
    """Читает старые файлы .xls (Excel 97-2003) через xlrd."""
    import xlrd

    participants = []
    book = xlrd.open_workbook(file_path)

    for sheet in book.sheets():
        header_row = None
        name_col = None
        position_col = None
        org_col = None
        group_col = None

        for row_idx in range(min(10, sheet.nrows)):
            for col_idx in range(sheet.ncols):
                hn = _norm_header(sheet.cell_value(row_idx, col_idx))
                if not hn:
                    continue
                if "фамилия" in hn or "фио" in hn or "имя" in hn:
                    header_row = row_idx
                    name_col = col_idx
                elif "должность" in hn or "профессия" in hn:
                    position_col = col_idx
                elif "подразделение" in hn or "организация" in hn or "наименование" in hn or hn == "сп":
                    org_col = col_idx
                elif "группа" in hn:
                    group_col = col_idx
            if header_row is not None:
                break

        if header_row is None or name_col is None:
            continue

        for row_idx in range(header_row + 1, sheet.nrows):
            row = sheet.row_values(row_idx)
            if not row or name_col >= len(row):
                continue

            name = _cell_text(row[name_col])
            if not name or len(name) < 3:
                continue
            
                        # Пропускаем мусор (вторая шапка и т.п.) - берём только похожее на ФИО.
            if not _looks_like_person(name):
                continue

            position = _cell_text(row[position_col]) if position_col is not None and len(row) > position_col else ""
            org = _cell_text(row[org_col]) if org_col is not None and len(row) > org_col else ""
            group = _cell_text(row[group_col]) if group_col is not None and len(row) > group_col else ""

            participants.append({
                "name": name,
                "position": position,
                "organization": org,
                "group": group,
            })

    return participants


def extract_from_docx(file_path: str) -> list[dict]:
    """Извлекает список участников из Word-файла (.docx)."""
    import docx

    participants = []
    doc = docx.Document(file_path)

    # Ищем таблицы с данными
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        # Проверяем заголовки
        header = [cell.text.lower().strip() for cell in table.rows[0].cells]

        name_idx = None
        position_idx = None
        org_idx = None
        group_idx = None

        for i, h in enumerate(header):
            if "фамилия" in h or "фио" in h or "имя" in h:
                name_idx = i
            if "должность" in h or "профессия" in h:
                position_idx = i
            if "организация" in h or "наименование" in h or "подразделение" in h:
                org_idx = i
            if "группа" in h:
                group_idx = i

        if name_idx is None:
            continue

        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]

            if not cells[name_idx] or len(cells[name_idx]) < 3:
                continue

            name = re.sub(r'\s+', ' ', cells[name_idx])
            position = re.sub(r'\s+', ' ', cells[position_idx]) if position_idx is not None and len(cells) > position_idx else ""
            org = re.sub(r'\s+', ' ', cells[org_idx]) if org_idx is not None and len(cells) > org_idx else ""
            group = re.sub(r'\s+', ' ', cells[group_idx]) if group_idx is not None and len(cells) > group_idx else ""

            participants.append({
                "name": name,
                "position": position,
                "organization": org,
                "group": group,
            })

    return participants


def extract_from_doc(file_path: str) -> list[dict]:
    """Конвертирует .doc в .docx и извлекает данные."""
    print("Конвертация .doc в .docx...")
    import win32com.client

    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False

    docx_path = file_path.replace(".doc", "_temp.docx")
    doc = word.Documents.Open(os.path.abspath(file_path))
    doc.SaveAs2(docx_path, FileFormat=16)
    doc.Close()
    word.Quit()

    try:
        result = extract_from_docx(docx_path)
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)

    return result


def manual_input_participants() -> list[dict]:
    """Ручной ввод участников (для PDF без OCR)."""
    print("\n=== РУЧНОЙ ВВОД УЧАСТНИКОВ ===")
    print("Введите данные участников. Для завершения введите пустое ФИО.\n")

    participants = []
    idx = 1

    while True:
        print(f"Участник {idx}:")
        name = input("  ФИО: ").strip()
        if not name:
            break

        position = input("  Должность: ").strip()
        org = input("  Организация: ").strip()

        participants.append({
            "name": name,
            "position": position,
            "organization": org,
        })
        idx += 1

    return participants


def extract_from_pdf(file_path: str) -> list[dict]:
    """
    Извлекает участников из скан-PDF заявки через локальный OCR
    (pdf_table_ocr.py, Tesseract). Работает полностью офлайн.

    Если в заявке несколько таблиц (например, отдельно по группам 1и2 и 3),
    пользователю предлагается выбрать нужную таблицу либо объединить все.
    """
    from pdf_table_ocr import extract_tables_from_pdf

    print("\nРаспознаю таблицы в PDF (офлайн, Tesseract)...")
    tables = extract_tables_from_pdf(file_path)

    if not tables:
        print("Таблицы не найдены или не распознаны.")
        print("Проверьте, что установлен русский языковой пакет Tesseract (tesseract-ocr-rus).")
        choice = input("Ввести данные вручную? (y/n): ").strip().lower()
        return manual_input_participants() if choice == "y" else []

    print(f"\nНайдено таблиц: {len(tables)}")
    for i, t in enumerate(tables, 1):
        hint = f" (похоже на группу {t['group_guess']})" if t["group_guess"] else ""
        print(f"  {i}. {t['heading'][:70] or '(без заголовка)'}{hint} — участников: {len(t['participants'])}")

    if len(tables) == 1:
        return tables[0]["participants"]

    print("\nВыберите номер таблицы для формирования протокола")
    print("(или 0, чтобы объединить всех участников из всех таблиц):")
    choice = input("Номер: ").strip()

    if choice == "0":
        combined = []
        for t in tables:
            combined.extend(t["participants"])
        return combined

    try:
        idx = int(choice) - 1
        return tables[idx]["participants"]
    except (ValueError, IndexError):
        print("Некорректный выбор, беру первую таблицу.")
        return tables[0]["participants"]


def extract_participants(file_path: str) -> list[dict]:
    """Главный метод извлечения - определяет формат и вызывает нужную функцию."""
    ext = Path(file_path).suffix.lower()

    if ext in (".xlsx", ".xls"):
        return extract_from_excel(file_path)
    elif ext == ".docx":
        return extract_from_docx(file_path)
    elif ext == ".doc":
        return extract_from_doc(file_path)
    elif ext == ".pdf":
        try:
            return extract_from_pdf(file_path)
        except ImportError as e:
            print(f"\nНе установлены зависимости для OCR: {e}")
            print("Установите: pip install pdf2image pytesseract opencv-python-headless numpy")
            print("И Tesseract OCR с русским языком (см. pdf_table_ocr.py).\n")
            choice = input("Ввести данные вручную? (y/n): ").strip().lower()
            return manual_input_participants() if choice == "y" else []
    else:
        raise ValueError(f"Неподдерживаемый формат: {ext}")


# ============================================================
# Генерация протокола
# ============================================================

def get_protocol_info() -> dict:
    """Запрашивает у пользователя данные для протокола."""
    print("\n=== ДАННЫЕ ДЛЯ ПРОТОКОЛА ===\n")

    info = {}

    info["protocol_number"] = input("Номер протокола: ").strip()
    info["date"] = input("Дата (ДД.ММ.ГГГГ) или Enter для сегодняшней: ").strip()
    if not info["date"]:
        info["date"] = datetime.now().strftime("%d.%m.%Y")

    info["course_name"] = input("Название курса: ").strip()
    if not info["course_name"]:
        info["course_name"] = "Безопасные методы и приемы выполнения работ на высоте"

    info["group"] = input("Группа (1, 2 или 3): ").strip()
    info["hours"] = input("Количество часов: ").strip()

    return info


def get_permit_codes_for_all(participants: list[dict]) -> dict:
    """Запрашивает коды допусков для каждого участника."""
    print("\n=== ДОПУСКИ ===")
    print("Введите коды допусков через запятую или диапазоны (например: 10,12-15)")
    print("Для справки введите 'help'\n")

    print(list_permits())
    print()

    permits_map = {}

    for participant in participants:
        name = participant["name"]
        print(f"\nДопуски для: {name}")

        while True:
            codes = input("  Коды допусков: ").strip()
            if codes.lower() == "help":
                print(list_permits())
                continue
            break

        if codes:
            permits_map[name] = parse_permits(codes)
            print(f"  → {permits_map[name][:100]}...")
        else:
            permits_map[name] = ""

    return permits_map


def create_protocol_docx(participants: list[dict], protocol_info: dict, permits_map: dict, output_path: str):
    """Создаёт Word-документ протокола."""
    import docx
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = docx.Document()

    # Настройка страницы
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    # Шапка
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ЧУДПО  "Учебный комбинат "Мелиоратор"')
    run.font.size = Pt(14)
    run.bold = True

    # Заголовок протокола
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'П Р О Т О К О Л    № {protocol_info["protocol_number"]}')
    run.font.size = Pt(14)
    run.bold = True

    # Дата
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'от  «{protocol_info["date"]}» г.')
    run.font.size = Pt(12)

    # Тип комиссии
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("заседания экзаменационной комиссии")
    run.font.size = Pt(12)

    # Состав комиссии
    doc.add_paragraph("в составе:")
    doc.add_paragraph("председатель комиссии: директор ЧУДПО УК «Мелиоратор» -  Первушов Александр Алексеевич")
    doc.add_paragraph("члены комиссии: заместитель директора ЧУДПО УК «Мелиоратор» -  Украинцев Павел Алексеевич")
    doc.add_paragraph("                              специалист по охране труда ЧУДПО УК «Мелиоратор» – Леонтьева Лариса Петровна")

    # Описание курса
    course_desc = (
        f'по выпуску окончивших обучение по курсу:  «{protocol_info["course_name"]} '
        f'для работников {protocol_info["group"]} группы» '
        f'в соответствии с Правилами по охране труда при работе на высоте '
        f'(приказ Минтруда России от 16.11.2020г. № 782н) '
        f'Произведя проверку знаний обучившихся по программе в объеме {protocol_info["hours"]} часов '
        f'комиссия постановила:'
    )
    doc.add_paragraph(course_desc)

    # Подписи
    doc.add_paragraph("")
    doc.add_paragraph("Председатель комиссии: _______________ А.А. Первушов")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Члены комиссии: ________________П.А. Украинцев")
    run.bold = True
    doc.add_paragraph("")
    doc.add_paragraph("_________________Л.П. Леонтьева")

    # Ведомость
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(f'В Е Д О М О С Т Ь к протоколу № {protocol_info["protocol_number"]}')
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'«{protocol_info["course_name"]}»')
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Конец обучения    от "{protocol_info["date"]}" г.')

    # Таблица участников
    doc.add_paragraph("")

    # План строк: человек с несколькими группами получает по строке на группу.
    rows_plan = []
    for participant in participants:
        if group_override:
            sub_groups = [group_override]
        else:
            sub_groups = self._split_groups(participant.get("group", ""))
            if not sub_groups:
                sub_groups = [protocol_info.get("group", "")]
        rows_plan.append((participant, sub_groups))

    total_rows = sum(len(sg) for _, sg in rows_plan)

    table = doc.add_table(rows=1 + total_rows, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "№\nп/п",
        "фамилия, имя, отчество,\nорганизация,",
        "Профессия\nдолжность",
        "заключение комиссии\nсдано/\nне сдано",
        "№ удосто-верения",
        "Группа\nпо безопас-\nности работ\nна высоте",
        "Допущен к работам:"
    ]

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.bold = True

    row_idx = 1
    for idx, (participant, sub_groups) in enumerate(rows_plan, 1):
        first_row = row_idx

        for gi, grp in enumerate(sub_groups):
            row = table.rows[row_idx]

            # ФИО/должность пишем один раз (потом объединим ячейки).
            if gi == 0:
                row.cells[0].text = str(idx)

                name_org = participant.get("name", "")
                if org_name:
                    name_org += "\n\n" + org_name
                row.cells[1].text = name_org

                pos = participant.get("position", "")
                sub = (participant.get("subdivision") or "").strip()
                if sub:
                    pos = (pos + "\n" + sub) if pos else sub
                row.cells[2].text = pos

            row.cells[3].text = "сдано"
            row.cells[4].text = "____"
            row.cells[5].text = grp

            # Допуски - только своей группы.
            texts = participant.get("permit_text_by_group", {}) or {}
            full_permits = texts.get(grp) or texts.get("") or ""
            if not full_permits:
                full_permits = participant.get("permit_text", "")
            row.cells[6].text = full_permits

            row_idx += 1

        # Несколько групп - объединяем №, ФИО и должность в общие ячейки.
        if len(sub_groups) > 1:
            for j in (0, 1, 2):
                top_cell = table.cell(first_row, j)
                for r in range(first_row + 1, row_idx):
                    top_cell = top_cell.merge(table.cell(r, j))

    # Форматирование всех ячеек.
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.save(output_path)
    print(f"\n✓ Протокол сохранён: {output_path}")

def _find_header_cols(row):
    """Ищет в строке шапку таблицы. Возвращает колонки или None, если это не шапка."""
    name_col = position_col = org_col = group_col = None
    for col_idx, cell in enumerate(row):
        hn = _norm_header(cell)
        if not hn:
            continue
        if "фамилия" in hn or "фио" in hn or "имя" in hn:
            name_col = col_idx
        elif "должность" in hn or "профессия" in hn:
            position_col = col_idx
        elif "подразделение" in hn or "организация" in hn or "наименование" in hn or hn == "сп":
            org_col = col_idx
        elif "группа" in hn:
            group_col = col_idx
    if name_col is None:
        return None
    return name_col, position_col, org_col, group_col


def _split_rows_into_tables(rows, label):
    """Разбивает строки листа на отдельные таблицы: как только встречается
    новая шапка с ФИО - начинается новая таблица."""
    tables = []
    current = None
    for row in rows:
        cols = _find_header_cols(row)
        if cols is not None:
            current = {
                "heading": f"{label} — таблица {len(tables) + 1}",
                "cols": cols,
                "participants": [],
            }
            tables.append(current)
            continue
        if current is None:
            continue
        name_col, position_col, org_col, group_col = current["cols"]
        if not row or name_col >= len(row):
            continue
        name = _cell_text(row[name_col])
        if not name or len(name) < 3 or not _looks_like_person(name):
            continue
        position = _cell_text(row[position_col]) if position_col is not None and len(row) > position_col else ""
        org = _cell_text(row[org_col]) if org_col is not None and len(row) > org_col else ""
        group = _cell_text(row[group_col]) if group_col is not None and len(row) > group_col else ""
        current["participants"].append({
            "name": name,
            "position": position,
            "organization": org,
            "group": group,
        })
    return [t for t in tables if t["participants"]]


def extract_excel_tables(file_path: str) -> list:
    """Находит в Excel-файле все таблицы (по листам, включая напечатанные
    друг под другом). Возвращает список, как у PDF: [{"heading", "participants"}]."""
    ext = Path(file_path).suffix.lower()
    tables = []
    if ext == ".xls":
        import xlrd
        book = xlrd.open_workbook(file_path)
        for sheet in book.sheets():
            rows = [sheet.row_values(r) for r in range(sheet.nrows)]
            tables.extend(_split_rows_into_tables(rows, sheet.name))
        return tables
    import openpyxl
    wb = openpyxl.load_workbook(file_path, data_only=True)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        tables.extend(_split_rows_into_tables(rows, sheet_name))
    return tables

# ============================================================
# Главная функция
# ============================================================

def main():
    if len(sys.argv) < 2:
        print(__doc__)
        print("\n=== СПРАВОЧНИК ДОПУСКОВ ===")
        print(list_permits())
        return

    input_file = sys.argv[1]

    if not os.path.exists(input_file):
        print(f"Ошибка: Файл не найден: {input_file}")
        return

    print(f"Обработка файла: {input_file}")

    # Извлекаем участников
    try:
        participants = extract_participants(input_file)
    except Exception as e:
        print(f"Ошибка извлечения данных: {e}")
        return

    if not participants:
        print("Не удалось извлечь участников из файла")
        return

    print(f"\nНайдено участников: {len(participants)}")
    for i, p in enumerate(participants, 1):
        print(f"  {i}. {p['name']} - {p['position']}")

    # Запрашиваем данные для протокола (можно передать через аргументы)
    if len(sys.argv) >= 6:
        protocol_info = {
            "protocol_number": sys.argv[2],
            "date": sys.argv[3] if sys.argv[3] else datetime.now().strftime("%d.%m.%Y"),
            "course_name": "Безопасные методы и приемы выполнения работ на высоте",
            "group": sys.argv[4],
            "hours": sys.argv[5],
        }
    else:
        protocol_info = get_protocol_info()

    # Запрашиваем допуски
    permits_map = get_permit_codes_for_all(participants)

    # Генерируем имя файла
    output_file = f"Протокол_{protocol_info['protocol_number']}_{protocol_info['date'].replace('.', '.')}.docx"

    # Создаём протокол
    create_protocol_docx(participants, protocol_info, permits_map, output_file)


if __name__ == "__main__":
    main()
